"""冒烟测试：建临时库 → 索引示例文件 → 中/英文搜索 → 状态与降级路径。
不依赖任何模型服务（图片应落在 waiting_model）。运行: uv run pytest -q
"""
from __future__ import annotations

import base64
import json
from pathlib import Path

import pytest

from semdex.config import Config
from semdex.db import Database
from semdex.indexer import index_pending
from semdex.models import ModelNotConfigured
from semdex.scanner import scan
from semdex.search import search
from semdex.textutil import build_fts_query, cjk_spaced, make_snippet

# 1x1 透明 PNG
PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
    "YPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)


@pytest.fixture()
def env(tmp_path: Path):
    data = tmp_path / "data"
    data.mkdir()
    (data / "会议记录.txt").write_text(
        "今天和张三讨论了广州地铁三号线的打卡方案，计划下周去体验东延段。", encoding="utf-8"
    )
    (data / "notes.md").write_text(
        "# Project Notes\nThe quick brown fox jumps over the lazy dog.\n", encoding="utf-8"
    )
    (data / "screenshot.png").write_bytes(PNG_1PX)
    (data / "unknown.zzz").write_bytes(b"\x00\x01binary")

    import docx
    d = docx.Document()
    d.add_paragraph("软考系统架构设计师复习计划")
    d.add_paragraph("重点：质量属性、架构风格、案例分析")
    d.save(str(data / "复习计划.docx"))

    cfg = Config(db_path=tmp_path / "index.db", folders=[data])
    db = Database(cfg.db_path)
    yield cfg, db, data
    db.close()


def _run_index(cfg, db):
    scan_stats = scan(db, cfg, log=lambda *_: None)
    index_stats = index_pending(db, cfg, log=lambda *_: None)
    return scan_stats, index_stats


def test_index_statuses(env):
    cfg, db, _ = env
    scan_stats, index_stats = _run_index(cfg, db)
    assert scan_stats.scanned == 5
    assert index_stats.indexed == 3          # txt + md + docx
    assert index_stats.waiting_model == 1    # png：视觉模型未启用
    assert index_stats.skipped == 1          # .zzz：无提取器
    by_status = db.counts()["by_status"]
    assert by_status.get("done") == 3
    assert by_status.get("waiting_model") == 1


def test_chinese_search(env):
    cfg, db, _ = env
    _run_index(cfg, db)
    # 两字词（trigram 方案的典型盲区，按字短语方案必须命中）
    hits = search(db, cfg, "地铁", mode="fulltext")
    assert hits and hits[0].filename == "会议记录.txt"
    assert "地铁" in hits[0].snippet
    # 人名、docx 内容
    assert search(db, cfg, "张三", mode="fulltext")
    hits = search(db, cfg, "软考", mode="fulltext")
    assert hits and hits[0].filename == "复习计划.docx"
    # 文件名命中
    assert search(db, cfg, "复习计划", mode="fulltext")


def test_english_and_hybrid_degrade(env):
    cfg, db, _ = env
    _run_index(cfg, db)
    hits = search(db, cfg, "quick fox", mode="fulltext")
    assert hits and hits[0].filename == "notes.md"
    # embedding 未启用：hybrid 自动退化为 fulltext，semantic 明确报错
    assert search(db, cfg, "quick", mode="hybrid")
    with pytest.raises(ModelNotConfigured):
        search(db, cfg, "quick", mode="semantic")


def test_incremental_and_delete(env):
    cfg, db, data = env
    _run_index(cfg, db)
    # 无变化重跑：不应产生新工作
    s2, i2 = _run_index(cfg, db)
    assert s2.new_or_changed == 0 and i2.indexed == 0
    # 修改文件 → 重新索引；删除文件 → 记录移除
    (data / "notes.md").write_text("totally new content about kubernetes", encoding="utf-8")
    (data / "会议记录.txt").unlink()
    s3, i3 = _run_index(cfg, db)
    assert s3.new_or_changed == 1 and s3.removed == 1 and i3.indexed == 1
    assert not search(db, cfg, "地铁", mode="fulltext")
    assert search(db, cfg, "kubernetes", mode="fulltext")


def test_hit_json_serializable(env):
    cfg, db, _ = env
    _run_index(cfg, db)
    hits = search(db, cfg, "地铁")
    payload = json.dumps([h.to_dict() for h in hits], ensure_ascii=False)
    assert "会议记录" in payload


def test_textutil():
    assert cjk_spaced("地铁2号线") .split() == ["地", "铁", "2", "号", "线"]
    assert build_fts_query("广州地铁 metro") == '"广 州 地 铁" "metro"'
    snippet = make_snippet("前面很长的内容" * 50 + "关键信息在这里" + "后面" * 50, "关键信息")
    assert "关键信息" in snippet
